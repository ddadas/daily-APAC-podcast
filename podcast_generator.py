"""
APAC Podcast Generator v18 - ENGAGING TWO-HOST EDITION
======================================================
Goal: a podcast that general, curious listeners genuinely enjoy and recommend.

What's new vs v17:
  1. TWO-HOST conversation format (Maya + Theo) instead of a single monologue.
  2. EDITORIAL CURATION pass: Claude ranks the day's stories for general-listener
     interest and we go DEEP on the best 3-4 instead of listing all 12.
  3. STORYTELLING prompt: humanized/rounded numbers, "why you should care" angle,
     hooks, and a recurring "one thing to tell a friend" segment.
  4. SHAREABILITY structure: cold-open hook -> deep dives -> tell-a-friend -> teaser.
  5. MECHANICAL fixes: dead-feed/empty-article filtering, de-duplication,
     Google-News RSS cleanup, number normalization for speech.
  6. DUAL-VOICE TTS: each host gets a distinct OpenAI voice, with a short beat
     of silence between turns so it breathes.

Still keeps factual grounding (no invented numbers/quotes/companies) but ALLOWS
plain-English explanatory context so a non-expert understands "why it matters".

Run it exactly like v17:  python Newspodcastv18_APAC_engaging.py
Requires in .env:  ANTHROPIC_API_KEY  and  OPENAI_API_KEY (for premium voices)
"""

import os, logging, datetime, re, asyncio, time, json, html, sys
from pathlib import Path
from datetime import datetime, timedelta, timezone

import requests
import feedparser

# Fix Windows console encoding for emoji support.
# line_buffering=True is CRITICAL: without it stdout is fully buffered and
# progress messages don't appear until the program ends (looks like "nothing
# is happening"). With it, each print() shows immediately.
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', line_buffering=True)
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', line_buffering=True)

import anthropic
from openai import OpenAI
from dotenv import load_dotenv
from pydub import AudioSegment, effects

# Optional article extraction
try:
    from newspaper import Article
    NEWSPAPER_AVAILABLE = True
except ImportError:
    NEWSPAPER_AVAILABLE = False
    print("WARNING: newspaper3k not installed. Install with: pip install newspaper3k")

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# Word document generation
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not installed. Install with: pip install python-docx")

EXTRACTED_ARTICLES = []

# ───────────────────────────────────────────────────────────────
# Hosts / voices  — tweak these to change the show's feel
# ───────────────────────────────────────────────────────────────
HOST_A_NAME = "Maya"      # warm anchor / storyteller
HOST_B_NAME = "Theo"      # curious co-host who asks the listener's questions
HOST_A_VOICE = "coral"    # gpt-4o-mini-tts voices: alloy, ash, ballad, coral, echo, fable, nova, onyx, sage, shimmer, verse
HOST_B_VOICE = "onyx"
SHOW_NAME = "The APAC Brief"

# TTS model. gpt-4o-mini-tts is newer, cheaper (~$0.015/min) and STEERABLE via the
# delivery instructions below. Falls back to tts-1-hd automatically if unavailable.
TTS_MODEL = "gpt-4o-mini-tts"
TTS_FALLBACK_MODEL = "tts-1-hd"

# Per-host delivery instructions — only gpt-4o-mini-tts honors these. This is what
# makes the two hosts sound like distinct people instead of two flat narrators.
HOST_INSTRUCTIONS = {
    HOST_A_NAME: ("Warm, engaging morning-radio host. Friendly and upbeat but credible. "
                  "Speak naturally and conversationally, with light energy and a smile in your voice. "
                  "Land the interesting moments; don't rush."),
    HOST_B_NAME: ("Curious, down-to-earth co-host. Genuinely interested, a little playful, sometimes "
                  "surprised. Ask questions like a smart friend who wants things explained simply. "
                  "Relaxed and natural, never stiff."),
}

# ───────────────────────────────────────────────────────────────
# Setup
# ───────────────────────────────────────────────────────────────
load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

# Quiet the noisy per-request HTTP logs from OpenAI/Anthropic/httpx so our own
# progress messages stay readable. (Errors still show.)
for _noisy in ("httpx", "openai", "anthropic", "urllib3", "newspaper"):
    logging.getLogger(_noisy).setLevel(logging.WARNING)

try:
    anthropic_client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
except Exception:
    print("ERROR: Could not initialize Anthropic client. Check ANTHROPIC_API_KEY in .env")
    sys.exit(1)

try:
    openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    OPENAI_AVAILABLE = bool(os.getenv("OPENAI_API_KEY"))
except Exception:
    OPENAI_AVAILABLE = False
    print("WARNING: OpenAI client not initialized - will use fallback TTS (single voice)")


# ───────────────────────────────────────────────────────────────
# News sourcing  (with dead-feed + dedup + cleanup)
# ───────────────────────────────────────────────────────────────
class APACNewsScraper:
    def __init__(self):
        self.feeds = {
            'The Straits Times': 'https://www.straitstimes.com/news/singapore/rss.xml',
            'Channel NewsAsia': 'https://www.channelnewsasia.com/api/v1/rss-outbound-feed?_format=xml',
            'Business Times Singapore': 'https://www.businesstimes.com.sg/rss-feeds-bt',
            'Economic Times': 'https://economictimes.indiatimes.com/rssfeedstopstories.cms',
            'Economic Times Markets': 'https://economictimes.indiatimes.com/markets/rssfeeds/1977021501.cms',
            'Business Standard': 'https://www.business-standard.com/rss/home_page_top_stories.rss',
            'The Hindu Business Line': 'https://www.thehindubusinessline.com/news/feeder/default.rss',
            'Mint': 'https://www.livemint.com/rss/markets',
            'Mint Economy': 'https://www.livemint.com/rss/economy',
            'Times of India Business': 'https://timesofindia.indiatimes.com/rssfeeds/1898055.cms',
            'Jakarta Post': 'https://www.thejakartapost.com/rss',
            'Bangkok Post Business': 'https://www.bangkokpost.com/rss/data/business.xml',
            'Vietnam News': 'https://vietnamnews.vn/rss/economy-and-politics.rss',
            'Philippine Star Business': 'https://www.philstar.com/rss/business',
            'Reuters Asia': 'https://news.google.com/rss/search?q=asia+business+when:1d&hl=en-SG&gl=SG&ceid=SG:en',
            'Nikkei Asia': 'https://news.google.com/rss/search?q=nikkei+asia+economy+when:1d&hl=en-SG&gl=SG&ceid=SG:en',
            'South China Morning Post': 'https://www.scmp.com/rss/91/feed',
            'Asia Times': 'https://asiatimes.com/feed/',
        }
        self.headers = {'User-Agent': 'Mozilla/5.0 (compatible; NewsBot/1.0; Educational Use)'}

    def get_recent_articles(self, hours=24, max_total_articles=20):
        cutoff_time = datetime.now() - timedelta(hours=hours)
        all_articles = []
        print("Fetching APAC news...")

        per_source = max(2, max_total_articles // max(1, len(self.feeds) // 2))

        for source_name, feed_url in self.feeds.items():
            try:
                articles = self._process_feed(source_name, feed_url, cutoff_time, per_source)
                all_articles.extend(articles)
                print(f"  OK  {source_name}: {len(articles)} usable articles")
            except Exception as e:
                print(f"  --  {source_name}: skipped ({e})")
            time.sleep(0.3)

        # De-duplicate by normalized title
        deduped = self._dedupe(all_articles)
        deduped.sort(key=lambda x: x['published'], reverse=True)
        final = deduped[:max_total_articles]
        print(f"Collected {len(final)} unique, non-empty articles (from {len(all_articles)} raw).")
        return final

    def _process_feed(self, source_name, feed_url, cutoff_time, max_articles):
        articles = []
        feed = feedparser.parse(feed_url)
        if not feed.entries:
            return articles

        processed = 0
        for entry in feed.entries:
            if processed >= max_articles:
                break

            title = getattr(entry, 'title', '').strip()
            # MECHANICAL FIX: skip blank-title / junk entries that broke v17
            if not title or len(title) < 8:
                continue

            # Date filter (keep undated items only if feed has no dates at all)
            pub_date = None
            if hasattr(entry, 'published_parsed') and entry.published_parsed:
                pub_date = datetime(*entry.published_parsed[:6])
                if pub_date < cutoff_time:
                    continue
            else:
                pub_date = datetime.now()

            article = self._extract_article(entry, source_name, title, pub_date)
            if article:
                articles.append(article)
                processed += 1
            time.sleep(0.2)
        return articles

    def _extract_article(self, entry, source_name, title, pub_date):
        try:
            raw_summary = getattr(entry, 'summary', '') or ''
            # Strip HTML tags that Google News / some feeds embed
            if BS4_AVAILABLE and ('<' in raw_summary):
                raw_summary = BeautifulSoup(raw_summary, 'html.parser').get_text(" ", strip=True)
            summary = html.unescape(raw_summary).strip()[:500]

            url = getattr(entry, 'link', '')
            body = ""
            if NEWSPAPER_AVAILABLE and url:
                try:
                    art = Article(url)
                    art.download()
                    art.parse()
                    if art.text:
                        body = art.text.strip()
                except Exception:
                    pass

            content = f"{summary}\n\n{body}".strip()

            # MECHANICAL FIX: drop articles with essentially no content
            if len(content.split()) < 25:
                return None

            return {
                'title': html.unescape(title),
                'source': source_name,
                'url': url,
                'published': pub_date,
                'summary': summary,
                'content_extract': content[:6000],  # cap to keep tokens sane
                'word_count': len(content.split()),
            }
        except Exception:
            return None

    @staticmethod
    def _dedupe(articles):
        seen = set()
        out = []
        for a in articles:
            key = a['title'].lower()
            key = re.sub(r'[^a-z0-9]', '', key)   # drop ALL spaces/punct so "top-10"=="top 10"
            key = key[:40]                         # first 40 chars as fingerprint
            if key in seen:
                continue
            seen.add(key)
            out.append(a)
        return out


# ───────────────────────────────────────────────────────────────
# STEP 1 — Editorial curation: pick the most interesting stories
# ───────────────────────────────────────────────────────────────
def curate_top_stories(articles, target=4):
    """Ask Claude to rank stories by general-listener interest and pick the best."""
    menu = ""
    for i, a in enumerate(articles, 1):
        menu += f"[{i}] {a['title']}  (source: {a['source']})\n"
        menu += f"     {a['summary'][:220]}\n\n"

    prompt = f"""You are the editor of a daily APAC news podcast for GENERAL, CURIOUS listeners
(not finance specialists). From the story menu below, choose the {target} stories that will be
MOST INTERESTING and ENJOYABLE for a broad audience and most likely to make someone say
"I have to tell my friend about this."

Score each story on: surprise/novelty, human stakes, "why should I care", and tell-a-friend factor.
Prefer a MIX of topics (not 4 market-recap stories). Avoid pure technical/corporate-action items
unless they have a genuinely interesting angle.

Return ONLY valid JSON in this exact shape, no prose:
{{"picks": [{{"id": <number>, "angle": "<one sentence: the human hook / why a normal person cares>"}}, ...]}}

STORY MENU:
{menu}"""

    try:
        resp = anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=800,
            messages=[{"role": "user", "content": prompt}],
        )
        text = "".join(b.text for b in resp.content if b.type == "text").strip()
        m = re.search(r'\{.*\}', text, re.DOTALL)
        data = json.loads(m.group(0))
        picks = []
        for p in data.get("picks", []):
            idx = int(p["id"]) - 1
            if 0 <= idx < len(articles):
                art = dict(articles[idx])
                art["editor_angle"] = p.get("angle", "")
                picks.append(art)
        if picks:
            print(f"Editor picked {len(picks)} stories:")
            for a in picks:
                print(f"   - {a['title'][:70]}  =>  {a['editor_angle'][:70]}")
            return picks
    except Exception as e:
        print(f"WARNING: curation failed ({e}); falling back to most-recent stories.")

    # Fallback: just take the first N
    return articles[:target]


# ───────────────────────────────────────────────────────────────
# STEP 2 — Write the engaging two-host script
# ───────────────────────────────────────────────────────────────
def build_story_brief(picks):
    brief = f"TODAY'S CHOSEN STORIES ({datetime.now().strftime('%B %d, %Y')}):\n\n"
    for i, a in enumerate(picks, 1):
        brief += f"STORY {i}: {a['title']}\n"
        brief += f"Source: {a['source']}\n"
        if a.get("editor_angle"):
            brief += f"Why it's interesting: {a['editor_angle']}\n"
        brief += f"Facts available:\n{a['content_extract']}\n"
        brief += "=" * 70 + "\n\n"
    return brief


def generate_two_host_script(picks):
    brief = build_story_brief(picks)

    prompt = f"""Write a lively, genuinely ENJOYABLE two-host news podcast script for {SHOW_NAME},
a daily show about Asia-Pacific for GENERAL, CURIOUS listeners (think smart friends, not finance pros).
Date: {datetime.now().strftime('%A, %B %d, %Y')}.

THE TWO HOSTS:
- {HOST_A_NAME}: the warm anchor who tells the story and explains things clearly.
- {HOST_B_NAME}: the curious co-host who asks the questions a normal listener is thinking,
  reacts ("wait, really?"), and keeps it grounded and human.

OUTPUT FORMAT — THIS IS STRICT:
- Write the script as alternating lines, each starting with the speaker's name and a colon:
      {HOST_A_NAME}: ...
      {HOST_B_NAME}: ...
- Put the speaker label ONLY when the speaker actually CHANGES. Never write two
  "{HOST_A_NAME}:" lines in a row — if one host says several sentences, keep them on ONE line
  under a single label. The hosts must genuinely alternate back and forth.
- NOTHING else. No stage directions, no [music], no headers, no asterisks, no markdown.
- Keep individual turns conversational (1-4 sentences each). Let them actually talk to each
  other, react, and hand off.
- HARD LIMIT: the whole script must be NO MORE THAN 50 speaker turns total. Stay tight.

STORYTELLING RULES (this is what makes it enjoyable):
1. COLD OPEN HOOK: Start with {HOST_A_NAME} teasing the 2-3 best things in today's episode in one
   punchy sentence ("Today: why a conflict thousands of kilometres away just torched Indian bank
   stocks, and the payments giant going public without raising a single dollar."). Then a quick warm
   welcome. Do NOT start with "Good morning and welcome..." boilerplate.
2. LEAD WITH WHY IT MATTERS, not with data. For every story, the FIRST thing is the human angle /
   the surprise / why a normal person should care. Numbers come second, sparingly.
3. HUMANIZE NUMBERS. Round hard. Say "about 448 billion dollars — gone in a week" NOT
   "four point four eight lakh crore rupees". Convert lakh/crore to plain billions/millions of
   dollars when you can, or to "roughly". Never read more than ONE precise figure per story, and
   never read decimals out loud. If a list of ten companies' market caps is in the source, do NOT
   recite it — summarize the pattern ("banks got hit hardest") and pick the single most striking number.
4. EXPLAIN, don't assume. If a story needs background a non-expert lacks (e.g. why oil prices matter
   to India), have {HOST_B_NAME} ask and {HOST_A_NAME} explain in plain English. General-knowledge
   explanation is ALLOWED and encouraged.
5. ONE "TELL A FRIEND" MOMENT: somewhere in the episode, include a short segment where a host shares
   the single most surprising / fun / quotable fact of the day, explicitly framed as the thing to
   tell someone.
6. SMOOTH HANDOFFS between stories — a host pivots naturally, sometimes connecting stories.
7. CLOSE with a forward-looking teaser ("tomorrow we're watching X") and a warm, short sign-off.

FACTUAL GROUNDING (still important):
- Do NOT invent specific numbers, company names, dates, or quotes that aren't in the facts below.
- You MAY add plain-English explanatory context that is general knowledge (how IPOs work, why oil
  matters to importers) — clearly as explanation, not as invented specifics.
- If you're unsure of a precise figure, speak in approximate terms ("roughly", "about") rather than
  inventing precision.

LENGTH: aim for a tight, fast-moving 8-11 minute episode (about 1300-1900 words of dialogue).
Better to be punchy than exhaustive. Cover the {len(picks)} stories below with real personality.

TODAY'S MATERIAL:
{brief}

Now write the full two-host script, alternating {HOST_A_NAME}: and {HOST_B_NAME}: lines only."""

    print("Writing the two-host script...")
    resp = anthropic_client.messages.create(
        model="claude-sonnet-4-5-20250929",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )
    script = "".join(b.text for b in resp.content if b.type == "text").strip()
    print(f"Script written: {len(script.split())} words.")
    return script


# ───────────────────────────────────────────────────────────────
# Parse script into (speaker, text) turns for dual-voice TTS
# ───────────────────────────────────────────────────────────────
def parse_turns(script):
    turns = []
    pattern = re.compile(rf'^\s*({re.escape(HOST_A_NAME)}|{re.escape(HOST_B_NAME)})\s*:\s*(.*)$')
    current_speaker, current_text = None, []

    for line in script.splitlines():
        m = pattern.match(line)
        if m:
            if current_speaker and current_text:
                turns.append((current_speaker, " ".join(current_text).strip()))
            current_speaker = m.group(1)
            current_text = [m.group(2).strip()]
        else:
            if current_speaker and line.strip():
                current_text.append(line.strip())
    if current_speaker and current_text:
        turns.append((current_speaker, " ".join(current_text).strip()))

    # MERGE consecutive same-speaker turns into one. If the model puts a label on
    # every sentence (Maya: ... Maya: ... Maya: ...), this collapses them so we make
    # ONE TTS call per actual speaker change instead of one per sentence.
    merged = []
    for spk, txt in turns:
        txt = txt.strip()
        if not txt:
            continue
        if merged and merged[-1][0] == spk:
            merged[-1] = (spk, (merged[-1][1] + " " + txt).strip())
        else:
            merged.append((spk, txt))

    # Clean each turn for speech
    cleaned = [(spk, clean_text_for_speech(txt)) for spk, txt in merged if txt.strip()]
    return cleaned


def clean_text_for_speech(text):
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'#+ (.+)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]', r'\1', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


# ───────────────────────────────────────────────────────────────
# Audio — dual voice with a beat of silence between turns
# ───────────────────────────────────────────────────────────────
# Safety ceiling: a normal 9-11 min episode is ~30-55 turns. If we ever exceed
# this, something went wrong (model padded the script / mis-labeled it) and we
# should NOT silently fire hundreds of paid TTS calls.
MAX_TTS_TURNS = 70


async def create_audio(turns, output_file):
    # Guard rail against runaway TTS billing
    if len(turns) > MAX_TTS_TURNS:
        print(f"\n  !!  WARNING: script produced {len(turns)} turns "
              f"(expected ~30-55). That's unusually high and would make many TTS calls.")
        print("      This usually means the script came back longer/odder than asked.")
        # In CI / non-interactive runs (e.g. GitHub Actions) there is no terminal to
        # answer the prompt, so we ABORT audio rather than hang or burn credits.
        non_interactive = os.getenv("CI") or os.getenv("GITHUB_ACTIONS") or not sys.stdin.isatty()
        if non_interactive:
            print("      Non-interactive run detected — skipping audio for safety. "
                  "Inspect the saved script; if it's fine, raise MAX_TTS_TURNS.")
            return False
        answer = input(f"      Continue and generate audio for all {len(turns)} turns anyway? [y/N]: ").strip().lower()
        if answer != "y":
            print("      Skipped audio. The text script + Word doc were still saved so you can inspect them.")
            return False

    if OPENAI_AVAILABLE:
        # Try the primary model first; if it errors (e.g. no access to the new model),
        # retry the whole episode once with the fallback model before giving up on dual-voice.
        for tts_model in (TTS_MODEL, TTS_FALLBACK_MODEL):
            if await _record_with_openai(turns, output_file, tts_model):
                return True
            if tts_model != TTS_FALLBACK_MODEL:
                print(f"   Retrying with fallback model {TTS_FALLBACK_MODEL}...")

    return await _record_with_edge(turns, output_file)


async def _record_with_openai(turns, output_file, tts_model):
    if OPENAI_AVAILABLE:
        try:
            print(f"Generating dual-voice audio with {tts_model} — {len(turns)} turns to record...")
            voice_map = {HOST_A_NAME: HOST_A_VOICE, HOST_B_NAME: HOST_B_VOICE}
            gap = AudioSegment.silent(duration=350)  # 0.35s breath between turns
            segments = []
            # gpt-4o-mini-tts caps input at ~2000 tokens (~1200-1500 chars), so split
            # smaller than the old tts-1-hd 3800. Turns are short anyway; this is a safety net.
            chunk_chars = 1200 if tts_model.startswith("gpt-4o") else 3800

            for i, (speaker, text) in enumerate(turns, 1):
                if not text:
                    continue
                voice = voice_map.get(speaker, HOST_A_VOICE)
                instructions = HOST_INSTRUCTIONS.get(speaker, "")
                for chunk in split_for_tts(text, chunk_chars):
                    # gpt-4o-mini-tts accepts a steerable `instructions` arg; tts-1-hd does not.
                    kwargs = dict(model=tts_model, voice=voice, input=chunk)
                    if tts_model.startswith("gpt-4o") and instructions:
                        kwargs["instructions"] = instructions
                    else:
                        kwargs["speed"] = 1.0
                    resp = openai_client.audio.speech.create(**kwargs)
                    temp = output_file.parent / f"_tmp_{i}.mp3"
                    with open(temp, 'wb') as f:
                        f.write(resp.content)
                    segments.append(AudioSegment.from_mp3(str(temp)))
                    os.unlink(temp)
                segments.append(gap)
                pct = int(i / len(turns) * 100)
                bar = "#" * (pct // 5) + "-" * (20 - pct // 5)
                print(f"   [{bar}] {pct:3d}%  turn {i}/{len(turns)} ({speaker}) recorded")
                time.sleep(0.15)

            print("   stitching and normalizing audio (almost done)...")
            combined = segments[0]
            for s in segments[1:]:
                combined += s
            combined = effects.normalize(combined)
            combined.export(output_file, format="mp3", bitrate="192k")
            print(f"OK: dual-voice audio saved ({tts_model}: {HOST_A_VOICE} + {HOST_B_VOICE}).")
            return True
        except Exception as e:
            print(f"WARNING: {tts_model} TTS failed: {e}")
    return False


async def _record_with_edge(turns, output_file):
    # Last-resort fallback: flatten to one voice via Edge TTS (free, no OpenAI).
    try:
        import edge_tts
        print("Falling back to single-voice Edge TTS...")
        flat = " ".join(f"{t}" for _, t in turns)
        communicate = edge_tts.Communicate(flat, "en-US-JennyNeural")
        temp_wav = output_file.with_suffix('.wav')
        await communicate.save(str(temp_wav))
        audio = effects.normalize(AudioSegment.from_wav(str(temp_wav)))
        audio.export(output_file, format="mp3", bitrate="128k")
        if temp_wav.exists():
            os.unlink(temp_wav)
        print("OK: single-voice fallback audio saved (Edge TTS).")
        return True
    except Exception as e:
        print(f"ERROR: all TTS failed: {e}")
        return False


def split_for_tts(text, max_chars=3800):
    if len(text) <= max_chars:
        return [text]
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks, cur = [], ""
    for s in sentences:
        if len(cur) + len(s) > max_chars and cur:
            chunks.append(cur.strip())
            cur = s
        else:
            cur += " " + s if cur else s
    if cur:
        chunks.append(cur.strip())
    return chunks


# ───────────────────────────────────────────────────────────────
# Outputs: Word doc + sources file
# ───────────────────────────────────────────────────────────────
def save_script_as_word(script, turns, output_file):
    if not DOCX_AVAILABLE:
        return False
    try:
        doc = Document()
        t = doc.add_heading(f'{SHOW_NAME} — Script', 0)
        t.alignment = 1
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", style='Subtitle')
        words = len(script.split())
        doc.add_paragraph(f"Hosts: {HOST_A_NAME} & {HOST_B_NAME}  |  ~{words} words  |  ~{words/150:.0f} min")
        doc.add_paragraph()
        for speaker, text in turns:
            p = doc.add_paragraph()
            run = p.add_run(f"{speaker}: ")
            run.bold = True
            p.add_run(text)
        doc.save(str(output_file))
        print(f"OK: Word script saved ({output_file.name}).")
        return True
    except Exception as e:
        print(f"WARNING: Word doc failed: {e}")
        return False


def save_sources_file(picks, all_articles, sources_file):
    lines = [f"{SHOW_NAME} — SOURCES",
             "=" * 40,
             f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
             f"Featured stories: {len(picks)}  |  Scanned pool: {len(all_articles)} articles", ""]
    lines.append("FEATURED IN THIS EPISODE:")
    for i, a in enumerate(picks, 1):
        lines.append(f"{i}. {a['title']}")
        lines.append(f"   {a['source']} — {a['published'].strftime('%B %d, %Y')}")
        lines.append(f"   {a['url']}")
        if a.get("editor_angle"):
            lines.append(f"   Angle: {a['editor_angle']}")
        lines.append("")
    Path(sources_file).write_text("\n".join(lines), encoding='utf-8')
    print(f"OK: sources file saved ({Path(sources_file).name}).")


# ───────────────────────────────────────────────────────────────
# Publishing — GitHub Releases as audio host + self-hosted RSS feed
# ───────────────────────────────────────────────────────────────
# Each episode becomes its own GitHub Release. The MP3 is uploaded as a release
# asset (free, public, no upload caps). A small episode.json sidecar carries
# duration + sources so we can rebuild the podcast RSS feed later without
# re-downloading the MP3. After each upload we regenerate docs/podcast.rss by
# walking the full list of releases — that file is served by GitHub Pages and is
# what Spotify (and Apple, Overcast, etc.) polls.

# Show metadata — override via env vars if needed.
PODCAST_TITLE = os.getenv("PODCAST_TITLE", "Your Daily APAC Business and Strategy Briefing")
PODCAST_AUTHOR = os.getenv("PODCAST_AUTHOR", "Daniel Burmester")
PODCAST_OWNER_EMAIL = os.getenv("PODCAST_OWNER_EMAIL", "burmesterdaniel6@gmail.com")
PODCAST_DESCRIPTION = os.getenv(
    "PODCAST_DESCRIPTION",
    "A daily two-host briefing on the most interesting business and economy stories "
    "across Asia-Pacific — Singapore, Southeast Asia, India and beyond. "
    f"Hosted by {HOST_A_NAME} and {HOST_B_NAME}.",
)
PODCAST_LANGUAGE = os.getenv("PODCAST_LANGUAGE", "en-us")
PODCAST_CATEGORY = os.getenv("PODCAST_CATEGORY", "Business")
PODCAST_SUBCATEGORY = os.getenv("PODCAST_SUBCATEGORY", "News")


def _pages_base_url():
    """The GitHub Pages base where docs/ is served from."""
    repo = os.getenv("GITHUB_REPOSITORY", "")  # "owner/repo"
    if "/" in repo:
        owner, name = repo.split("/", 1)
        return f"https://{owner}.github.io/{name}"
    return os.getenv("PODCAST_SITE_URL", "https://example.github.io/daily-APAC-podcast")


def _cover_url():
    return os.getenv("PODCAST_COVER_URL", f"{_pages_base_url()}/cover.jpg")


def build_episode_description(picks):
    """Short HTML show-notes built from the featured stories + sources."""
    parts = [f"<p>Today on {SHOW_NAME}, {HOST_A_NAME} and {HOST_B_NAME} "
             "break down the most interesting business and economy stories across Asia-Pacific.</p>",
             "<p><strong>In this episode:</strong></p><ul>"]
    for a in picks:
        parts.append(f"<li>{html.escape(a['title'])}</li>")
    parts.append("</ul><p><strong>Sources:</strong></p><ul>")
    for a in picks:
        parts.append(f'<li><a href="{html.escape(a["url"])}">{html.escape(a["source"])}</a></li>')
    parts.append("</ul>")
    return "".join(parts)


def build_release_body_markdown(picks):
    """Markdown shown on the GitHub Release page itself (human-readable)."""
    lines = [f"Daily {SHOW_NAME} episode with {HOST_A_NAME} and {HOST_B_NAME}.", "",
             "**In this episode:**"]
    for a in picks:
        lines.append(f"- {a['title']}")
    lines.append("")
    lines.append("**Sources:**")
    for a in picks:
        lines.append(f"- [{a['source']}]({a['url']}) — {a['title']}")
    return "\n".join(lines)


def _get_audio_duration_seconds(mp3_path):
    try:
        audio = AudioSegment.from_mp3(str(mp3_path))
        return int(round(len(audio) / 1000))
    except Exception:
        return 0


def _format_duration_hms(seconds):
    h, rem = divmod(int(seconds), 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _rfc822(dt):
    """RFC 822 / RFC 1123 date for RSS pubDate. dt must be timezone-aware."""
    if dt.tzinfo is None:
        dt = dt.astimezone()
    return dt.strftime("%a, %d %b %Y %H:%M:%S %z")


def _gh_headers(token):
    return {
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
        "User-Agent": f"{SHOW_NAME}-generator/1.0",
    }


def publish_to_github_releases(mp3_file, picks, episode_title):
    """Create a GitHub Release for this episode and rebuild docs/podcast.rss.

    Requires the following env vars (auto-set in GitHub Actions):
      GITHUB_TOKEN       - token with `contents: write` on this repo
      GITHUB_REPOSITORY  - owner/repo, e.g. "ddadas/daily-APAC-podcast"

    Returns the release dict on success, None on skip/failure.
    """
    token = os.getenv("GITHUB_TOKEN")
    repo = os.getenv("GITHUB_REPOSITORY")
    if not (token and repo):
        print("GitHub Release: skipped (GITHUB_TOKEN / GITHUB_REPOSITORY not set).")
        print("    To publish locally, run inside Actions or set both env vars manually.")
        return None

    tag = Path(mp3_file).stem  # e.g. "apac_brief_2026-06-16_05-00"
    duration_seconds = _get_audio_duration_seconds(mp3_file)
    size_bytes = Path(mp3_file).stat().st_size
    pub_at = datetime.now().astimezone()

    # Create the release
    api_root = f"https://api.github.com/repos/{repo}/releases"
    headers = _gh_headers(token)
    create_payload = {
        "tag_name": tag,
        "name": episode_title,
        "body": build_release_body_markdown(picks),
        "draft": False,
        "prerelease": False,
    }
    print(f"GitHub Release: creating release '{tag}'...")
    resp = requests.post(api_root, json=create_payload, headers=headers, timeout=60)
    if resp.status_code not in (200, 201):
        print(f"GitHub Release: create failed (HTTP {resp.status_code}): {resp.text[:300]}")
        return None
    release = resp.json()
    upload_base = release["upload_url"].split("{")[0]

    # Upload the MP3
    print(f"GitHub Release: uploading MP3 ({size_bytes / 1024 / 1024:.1f} MB)...")
    mp3_name = Path(mp3_file).name
    with open(mp3_file, "rb") as f:
        mp3_bytes = f.read()
    up = requests.post(
        upload_base,
        params={"name": mp3_name, "label": "Episode audio (MP3)"},
        headers={**headers, "Content-Type": "audio/mpeg"},
        data=mp3_bytes,
        timeout=600,
    )
    if up.status_code not in (200, 201):
        print(f"GitHub Release: MP3 upload failed (HTTP {up.status_code}): {up.text[:300]}")
        return None
    mp3_asset = up.json()

    # Upload episode.json sidecar so future RSS rebuilds can recover the metadata
    episode_meta = {
        "title": episode_title,
        "tag": tag,
        "duration_seconds": duration_seconds,
        "size_bytes": size_bytes,
        "published_at": pub_at.isoformat(),
        "description_html": build_episode_description(picks),
        "summary": f"Daily {SHOW_NAME} briefing with {HOST_A_NAME} and {HOST_B_NAME}.",
        "stories": [
            {"title": a["title"], "source": a["source"], "url": a["url"],
             "angle": a.get("editor_angle", "")}
            for a in picks
        ],
    }
    meta_bytes = json.dumps(episode_meta, indent=2).encode("utf-8")
    print("GitHub Release: uploading episode.json sidecar...")
    up2 = requests.post(
        upload_base,
        params={"name": "episode.json", "label": "Episode metadata"},
        headers={**headers, "Content-Type": "application/json"},
        data=meta_bytes,
        timeout=60,
    )
    if up2.status_code not in (200, 201):
        # Non-fatal: RSS rebuild can fall back to release body
        print(f"GitHub Release: episode.json upload failed (HTTP {up2.status_code}). Continuing.")

    # Re-fetch the release so we get the fully-populated asset list. The list
    # endpoint used later for RSS regen reads from a cache that briefly lags
    # behind asset uploads, so we hand the regen the canonical record directly.
    refreshed = requests.get(
        f"{api_root}/{release['id']}",
        headers=headers,
        timeout=30,
    )
    if refreshed.status_code == 200:
        release = refreshed.json()

    print(f"OK: GitHub Release published — {release['html_url']}")
    print(f"    MP3 URL: {mp3_asset['browser_download_url']}")
    return release


def regenerate_podcast_rss(rss_path, ensure_release=None):
    """Rebuild the podcast RSS feed by walking all releases in the repo.

    The RSS feed lives at docs/podcast.rss and is served by GitHub Pages.
    Spotify / Apple / Overcast all poll it.

    `ensure_release`, if provided, is a release dict that MUST be present in
    the output even if GitHub's list endpoint hasn't surfaced it yet (it has
    a short cache lag right after creation).
    """
    token = os.getenv("GITHUB_TOKEN")
    repo = os.getenv("GITHUB_REPOSITORY")
    if not repo:
        print("RSS regen: skipped (GITHUB_REPOSITORY not set).")
        return False

    headers = _gh_headers(token) if token else {"Accept": "application/vnd.github+json"}

    releases = []
    page = 1
    while True:
        r = requests.get(
            f"https://api.github.com/repos/{repo}/releases",
            params={"per_page": 100, "page": page},
            headers=headers,
            timeout=60,
        )
        if r.status_code != 200:
            print(f"RSS regen: failed to list releases (HTTP {r.status_code}): {r.text[:200]}")
            return False
        batch = r.json()
        if not batch:
            break
        releases.extend(batch)
        if len(batch) < 100:
            break
        page += 1

    # Merge in the just-created release if the list endpoint hasn't caught up.
    if ensure_release:
        if not any(r.get("id") == ensure_release.get("id") for r in releases):
            releases.insert(0, ensure_release)

    items_xml = []
    kept = 0
    for rel in releases:
        if rel.get("draft") or rel.get("prerelease"):
            continue
        mp3_asset = next(
            (a for a in rel.get("assets", []) if a["name"].lower().endswith(".mp3")),
            None,
        )
        if not mp3_asset:
            continue
        meta_asset = next(
            (a for a in rel.get("assets", []) if a["name"] == "episode.json"),
            None,
        )
        # Pull metadata from sidecar if available
        duration_seconds = 0
        description_html = rel.get("body") or ""
        if meta_asset:
            try:
                # Asset is public for public repos — no auth needed.
                m = requests.get(meta_asset["browser_download_url"], timeout=30)
                if m.status_code == 200:
                    meta = m.json()
                    duration_seconds = int(meta.get("duration_seconds", 0))
                    description_html = meta.get("description_html") or description_html
            except Exception:
                pass

        pub_iso = rel.get("published_at") or rel.get("created_at")
        try:
            pub_dt = datetime.strptime(pub_iso, "%Y-%m-%dT%H:%M:%SZ").replace(tzinfo=timezone.utc)
        except Exception:
            pub_dt = datetime.now().astimezone()

        title = rel.get("name") or rel.get("tag_name")
        guid = rel.get("tag_name")
        mp3_url = mp3_asset["browser_download_url"]
        mp3_size = mp3_asset.get("size", 0)

        items_xml.append(_render_item(
            title=title,
            guid=guid,
            url=mp3_url,
            size_bytes=mp3_size,
            duration_seconds=duration_seconds,
            pub_dt=pub_dt,
            description_html=description_html,
        ))
        kept += 1

    rss_xml = _render_channel(items_xml)
    Path(rss_path).parent.mkdir(parents=True, exist_ok=True)
    Path(rss_path).write_text(rss_xml, encoding="utf-8")
    print(f"OK: RSS feed regenerated with {kept} episode(s) -> {rss_path}")
    return True


def _render_item(title, guid, url, size_bytes, duration_seconds, pub_dt, description_html):
    safe_title = html.escape(title)
    safe_guid = html.escape(guid)
    duration_str = _format_duration_hms(duration_seconds) if duration_seconds else "10:00"
    return f"""    <item>
      <title>{safe_title}</title>
      <description><![CDATA[{description_html}]]></description>
      <content:encoded><![CDATA[{description_html}]]></content:encoded>
      <pubDate>{_rfc822(pub_dt)}</pubDate>
      <guid isPermaLink="false">{safe_guid}</guid>
      <enclosure url="{html.escape(url)}" length="{int(size_bytes)}" type="audio/mpeg" />
      <itunes:author>{html.escape(PODCAST_AUTHOR)}</itunes:author>
      <itunes:duration>{duration_str}</itunes:duration>
      <itunes:summary>{safe_title}</itunes:summary>
      <itunes:explicit>false</itunes:explicit>
    </item>"""


def _render_channel(items_xml):
    title = html.escape(PODCAST_TITLE)
    desc = html.escape(PODCAST_DESCRIPTION)
    author = html.escape(PODCAST_AUTHOR)
    email = html.escape(PODCAST_OWNER_EMAIL)
    cover = html.escape(_cover_url())
    site = html.escape(_pages_base_url())
    items_block = "\n".join(items_xml) if items_xml else ""
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<rss version="2.0"
     xmlns:itunes="http://www.itunes.com/dtds/podcast-1.0.dtd"
     xmlns:content="http://purl.org/rss/1.0/modules/content/"
     xmlns:atom="http://www.w3.org/2005/Atom">
  <channel>
    <title>{title}</title>
    <link>{site}</link>
    <atom:link href="{site}/podcast.rss" rel="self" type="application/rss+xml" />
    <description>{desc}</description>
    <language>{html.escape(PODCAST_LANGUAGE)}</language>
    <copyright>(c) {datetime.now().year} {author}</copyright>
    <lastBuildDate>{_rfc822(datetime.now().astimezone())}</lastBuildDate>
    <itunes:author>{author}</itunes:author>
    <itunes:summary>{desc}</itunes:summary>
    <itunes:owner>
      <itunes:name>{author}</itunes:name>
      <itunes:email>{email}</itunes:email>
    </itunes:owner>
    <itunes:image href="{cover}" />
    <itunes:explicit>false</itunes:explicit>
    <itunes:category text="{html.escape(PODCAST_CATEGORY)}">
      <itunes:category text="{html.escape(PODCAST_SUBCATEGORY)}" />
    </itunes:category>
    <itunes:type>episodic</itunes:type>
{items_block}
  </channel>
</rss>
"""


# ───────────────────────────────────────────────────────────────
# Main
# ───────────────────────────────────────────────────────────────
def main():
    global EXTRACTED_ARTICLES
    print("\n" + "=" * 70)
    print(f"  {SHOW_NAME} — Engaging Two-Host Generator (v18)")
    print(f"  Hosts: {HOST_A_NAME} ({HOST_A_VOICE}) + {HOST_B_NAME} ({HOST_B_VOICE})")
    print("=" * 70)

    if not os.getenv("ANTHROPIC_API_KEY"):
        print("ERROR: ANTHROPIC_API_KEY required in .env")
        return

    try:
        # Always write output NEXT TO THIS SCRIPT FILE, not next to wherever the
        # terminal happened to be when you launched it. (Previously used a relative
        # path, which scattered episodes into the current working directory.)
        script_dir = Path(__file__).resolve().parent
        out_dir = script_dir / "podcast_episodes"
        out_dir.mkdir(exist_ok=True)
        print(f"    Output folder: {out_dir}")
        stamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        base = f"apac_brief_{stamp}"
        mp3_file = out_dir / f"{base}.mp3"
        script_file = out_dir / f"{base}_script.txt"
        word_file = out_dir / f"{base}_script.docx"
        sources_file = out_dir / f"{base}_sources.txt"

        # Gather + clean
        print("\n[1/4] FETCHING: scanning APAC news sources...")
        scraper = APACNewsScraper()
        articles = scraper.get_recent_articles(hours=24, max_total_articles=20)
        if not articles:
            raise Exception("No usable articles found")
        EXTRACTED_ARTICLES = articles

        # Curate the best stories for a general audience
        print("\n[2/4] EDITOR: choosing the most interesting stories...")
        picks = curate_top_stories(articles, target=4)

        # Write the engaging two-host script
        print("\n[3/4] WRITER: composing the episode...")
        script = generate_two_host_script(picks)
        turns = parse_turns(script)
        if not turns:
            raise Exception("Script could not be parsed into host turns — check format.")

        # EARLY HEALTH CHECK: show turn count + rough cost BEFORE any paid audio.
        # Healthy episode = ~30-55 turns. Way more than that = something went wrong.
        n_turns = len(turns)
        n_words = len(script.split())
        if n_turns <= MAX_TTS_TURNS:
            health = "looks healthy"
        else:
            health = "UNUSUALLY HIGH — will prompt before recording"
        print(f"\n    Script ready: {n_words} words across {n_turns} speaker turns ({health}).")
        print(f"    That means about {n_turns} text-to-speech calls and roughly a {n_words/150:.0f}-minute episode.")

        # Save text + word + sources
        script_file.write_text(script, encoding='utf-8')
        print(f"OK: text script saved ({script_file.name}).")
        save_script_as_word(script, turns, word_file)
        save_sources_file(picks, articles, sources_file)

        # Audio
        print(f"\n[4/4] NARRATOR: recording the episode ({n_turns} turns)...")
        print("    (each line below = one turn recorded; this is the slow part)")
        ok = asyncio.run(create_audio(turns, mp3_file))

        if ok:
            size_mb = mp3_file.stat().st_size / (1024 * 1024)
            words = len(script.split())
            print("\n" + "=" * 70)
            print("DONE — episode generated!")
            print(f"  Audio:   {mp3_file.absolute()}")
            print(f"  Script:  {script_file.absolute()}")
            print(f"  Word:    {word_file.absolute()}")
            print(f"  Sources: {sources_file.absolute()}")
            print(f"  {words} words | ~{words/150:.0f} min | {size_mb:.1f} MB | {len(turns)} turns")
            print("=" * 70)

            # Publish to GitHub Releases and regenerate docs/podcast.rss.
            # Spotify polls the RSS feed served by GitHub Pages.
            episode_title = f"{PODCAST_TITLE} — {datetime.now().strftime('%A, %B %d, %Y')}"
            print("\n[5/5] PUBLISHING to GitHub Releases + RSS:")
            release = publish_to_github_releases(mp3_file, picks, episode_title)
            if release:
                rss_path = script_dir / "docs" / "podcast.rss"
                regenerate_podcast_rss(rss_path, ensure_release=release)
                print(f"\n    Feed URL (submit this once to Spotify for Podcasters):")
                print(f"    {_pages_base_url()}/podcast.rss")
        else:
            print("\nERROR: audio generation failed (script + docs were still saved).")

    except Exception as e:
        print(f"\nERROR: {e}")
        logging.error(f"v18 generation failed: {e}")


if __name__ == "__main__":
    main()
